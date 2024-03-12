import sys

sys.path.append("../src")
from docxTools import in_which_run_is, at_which_position_in_its_run_is, cp, rm, mv
from docx import Document

def test_docxTools():
    document = Document()

    p_src = document.add_paragraph("")

    p_src.add_run("abc").bold = True
    p_src.add_run("defg").underline = True
    p_src.add_run("hijkl").italic = True

    assert in_which_run_is(0, p_src) == 0
    assert in_which_run_is(1, p_src) == 0
    assert in_which_run_is(2, p_src) == 0

    assert in_which_run_is(3, p_src) == 1
    assert in_which_run_is(4, p_src) == 1
    assert in_which_run_is(5, p_src) == 1
    assert in_which_run_is(6, p_src) == 1

    assert in_which_run_is(7, p_src) == 2
    assert in_which_run_is(8, p_src) == 2
    assert in_which_run_is(9, p_src) == 2
    assert in_which_run_is(10, p_src) == 2
    assert in_which_run_is(11, p_src) == 2

    assert in_which_run_is(12, p_src) == None
    assert in_which_run_is(-1, p_src) == None

    assert at_which_position_in_its_run_is(0, p_src) == 0
    assert at_which_position_in_its_run_is(1, p_src) == 1
    assert at_which_position_in_its_run_is(2, p_src) == 2

    assert at_which_position_in_its_run_is(3, p_src) == 0
    assert at_which_position_in_its_run_is(4, p_src) == 1
    assert at_which_position_in_its_run_is(5, p_src) == 2
    assert at_which_position_in_its_run_is(6, p_src) == 3

    assert at_which_position_in_its_run_is(7, p_src) == 0
    assert at_which_position_in_its_run_is(8, p_src) == 1
    assert at_which_position_in_its_run_is(9, p_src) == 2
    assert at_which_position_in_its_run_is(10, p_src) == 3
    assert at_which_position_in_its_run_is(11, p_src) == 4

    assert at_which_position_in_its_run_is(12, p_src) == None
    assert at_which_position_in_its_run_is(-1, p_src) == None

    p_dest = document.add_paragraph("")

    cp(5, 7, p_src, p_dest)
    assert p_dest.text == "fgh"
    p_dest._p.clear()

    cp(0, 0, p_src, p_dest)
    assert p_dest.text == "a"
    p_dest._p.clear()

    cp(0, 2, p_src, p_dest)
    assert p_dest.text == "abc"
    p_dest._p.clear()

    cp(0, 3, p_src, p_dest)
    assert p_dest.text == "abcd"
    p_dest._p.clear()

    cp(0, 11, p_src, p_dest)
    assert p_dest.text == "abcdefghijkl"
    p_dest._p.clear()

    cp(7, 11, p_src, p_dest)
    assert p_dest.text == "hijkl"
    p_dest._p.clear()

    cp(11, 11, p_src, p_dest)
    assert p_dest.text == "l"
    p_dest._p.clear()

    cp(0, 11, p_src, p_dest)

    rm(0, 0, p_dest)
    assert p_dest.text == "bcdefghijkl"

    p_dest._p.clear()
    cp(0, 11, p_src, p_dest)
    rm(5, 7, p_dest)
    assert p_dest.text == "abcdeijkl"

    p_dest._p.clear()
    cp(0, 11, p_src, p_dest)
    rm(0, 2, p_dest)
    assert p_dest.text == "defghijkl"

    p_dest._p.clear()
    cp(0, 11, p_src, p_dest)
    rm(0, 11, p_dest)
    assert p_dest.text == ""

    p_dest._p.clear()
    cp(0, 11, p_src, p_dest)
    rm(1, 1, p_dest)
    assert p_dest.text == "acdefghijkl"

    p_dest._p.clear()
    cp(0, 11, p_src, p_dest)
    rm(1, 12, p_dest)
    assert p_dest.text == "abcdefghijkl"

    p_dest._p.clear()
    cp(0, 11, p_src, p_dest)
    rm(0, 2, p_dest)
    assert p_dest.text == "defghijkl"

    p_dest._p.clear()
    cp(0, 11, p_src, p_dest)
    rm(3, 6, p_dest)
    assert p_dest.text == "abchijkl"

    p_dest._p.clear()
    mv(0, 0, p_src, p_dest)
    assert p_src.text == "bcdefghijkl"
    assert p_dest.text == "a"

    p_src._p.clear()
    p_src.add_run("abc").bold = True
    p_src.add_run("defg").underline = True
    p_src.add_run("hijkl").italic = True
    p_dest._p.clear()
    mv(0, 11, p_src, p_dest)
    assert p_src.text == ""
    assert p_dest.text == "abcdefghijkl"

    p_src._p.clear()
    p_src.add_run("abc").bold = True
    p_src.add_run("defg").underline = True
    p_src.add_run("hijkl").italic = True
    p_dest._p.clear()
    mv(3, 6, p_src, p_dest)
    assert p_src.text == "abchijkl"
    assert p_dest.text == "defg"


    # document.save("./test.docx")
