from docx import Document
from docx_tools import extractTextBetween
import pytest


def test_extract_substring_between_in_empty_string_returns_string():
    doc = Document()
    doc.add_paragraph("")
    assert extractTextBetween(doc, 0, 0, 0, 5) == ""


def test_extract_substring_between_in_empty_string_returns_empty_string1():
    doc = Document()
    doc.add_paragraph("")
    assert extractTextBetween(doc, 0, 0, 0, 0) == ""


def test_extract_substring_between_in_empty_string_returns_empty_string2():
    doc = Document()
    doc.add_paragraph("")
    assert extractTextBetween(doc, 0, 0, 0, 1) == ""


def test_extract_substring_between_in_empty_string_returns_nonempty_string1():
    doc = Document()
    doc.add_paragraph("asd")
    assert extractTextBetween(doc, 0, 0, 0, 1) == "a"


def test_extract_substring_between_in_empty_string_returns_nonempty_string2():
    doc = Document()
    doc.add_paragraph("asd")
    assert extractTextBetween(doc, 0, 0, 0, 2) == "as"


def test_extract_substring_between_in_empty_string_returns_nonempty_string3():
    doc = Document()
    doc.add_paragraph("asd")
    assert extractTextBetween(doc, 0, 0, 0, 3) == "asd"


def test_extract_substring_between_in_empty_string_returns_nonempty_string4():
    doc = Document()
    doc.add_paragraph("asd")
    assert extractTextBetween(doc, 0, 0, 0, 5) == "asd"


def test_extract_substring_between_in_empty_string_returns_nonempty_string5():
    doc = Document()
    doc.add_paragraph("asd")
    assert extractTextBetween(doc, 0, 0, 1, 3) == "sd"


def test_extract_substring_between_in_empty_string_returns_nonempty_string6():
    doc = Document()
    doc.add_paragraph("asd")
    assert extractTextBetween(doc, 0, 0, 2, 3) == "d"


def test_extract_substring_between_in_empty_string_returns_nonempty_string7():
    doc = Document()
    doc.add_paragraph("asd")
    assert extractTextBetween(doc, 0, 0, 3, 3) == ""


def test_substring_betweeen_to_empty_paragraphs_is_empty():
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("")
    assert extractTextBetween(doc, 0, 1, 0, 0) == ""


def test_substring_betweeen_to_nonempty_paragraphs_is_empty():
    doc = Document()
    doc.add_paragraph("AUTO")
    doc.add_paragraph("BAUM")
    assert extractTextBetween(doc, 1, 1, 0, 0) == ""


def test_start_paragraph_is_out_of_bounds1():
    doc = Document()
    doc.add_paragraph("")
    with pytest.raises(Exception):
        extractTextBetween(doc, 0, -1, 0, 0)


def test_start_paragraph_is_out_of_bounds2():
    doc = Document()
    doc.add_paragraph("")
    with pytest.raises(Exception):
        extractTextBetween(doc, 3, 0, 0, 0)


def test_start_paragraph_is_out_of_bounds3():
    doc = Document()
    doc.add_paragraph("")
    with pytest.raises(Exception):
        extractTextBetween(doc, 0, -1, 0, 0)


def test_start_paragraph_is_out_of_bounds4():
    doc = Document()
    doc.add_paragraph("")
    with pytest.raises(Exception):
        extractTextBetween(doc, 0, 3, 0, 0)


def test_substring_betweeen_to_nonempty_paragraphs_is_nonempty1():
    doc = Document()
    doc.add_paragraph("AUTO")
    doc.add_paragraph("BAUM")
    assert extractTextBetween(doc, 1, 1, 0, 4) == "BAUM"


def test_substring_betweeen_to_nonempty_paragraphs_is_nonempty2():
    doc = Document()
    doc.add_paragraph("AUTO")
    doc.add_paragraph("BAUM")
    assert extractTextBetween(doc, 0, 1, 0, 0) == "AUTO"
