from docx_tools import copyTextSegment
from docx import Document

def test_cp():
    document = Document()

    p_src = document.add_paragraph("")

    p_src.add_run("abc").bold = True
    p_src.add_run("defg").underline = True
    p_src.add_run("hijkl").italic = True
    p_dest = document.add_paragraph("")

    copyTextSegment(p_src, p_dest, 5, 7)
    assert p_dest.text == "fgh"
    p_dest._p.clear()

    copyTextSegment(p_src, p_dest, 0, 0)
    assert p_dest.text == "a"
    p_dest._p.clear()

    copyTextSegment(p_src, p_dest, 0, 2)
    assert p_dest.text == "abc"
    p_dest._p.clear()

    copyTextSegment(p_src, p_dest, 0, 3)
    assert p_dest.text == "abcd"
    p_dest._p.clear()

    copyTextSegment(p_src, p_dest, 0, 11)
    assert p_dest.text == "abcdefghijkl"
    p_dest._p.clear()

    copyTextSegment(p_src, p_dest, 7, 11)
    assert p_dest.text == "hijkl"
    p_dest._p.clear()

    copyTextSegment(p_src, p_dest, 11, 11)
    assert p_dest.text == "l"
    p_dest._p.clear()

    copyTextSegment(p_src, p_dest, 0, 11)