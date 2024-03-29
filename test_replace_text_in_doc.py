from docx import Document
import docx_tools

def test_replace_text_in_doc1():
    doc = Document()
    doc.add_paragraph("")
    docx_tools.replaceDocTextSegment(doc, 0, 0, 0, 0, "FOO")
    assert docx_tools.extractOuterDocText(doc) == "FOO"


def test_replace_text_in_doc2():
    doc = Document()
    doc.add_paragraph("")
    docx_tools.replaceDocTextSegment(doc, 0, 0, 0, 0, "BAR")
    assert docx_tools.extractOuterDocText(doc) == "BAR"


def test_replace_text_in_doc3():
    doc = Document()
    doc.add_paragraph("FOO")
    docx_tools.replaceDocTextSegment(doc, 0, 0, 0, 2, "BAR")
    assert docx_tools.extractOuterDocText(doc) == "BAR"


def test_replace_text_in_doc4():
    doc = Document()
    doc.add_paragraph("FOO")
    doc.add_paragraph("BAR")
    docx_tools.replaceDocTextSegment(doc, 0, 1, 1, 2, "Hello")
    assert docx_tools.extractOuterDocText(doc) == "FHello"