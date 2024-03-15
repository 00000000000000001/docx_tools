from docx_tools import moveTextSegment
from docx import Document

def test_mv():
    document = Document()

    p_src = document.add_paragraph("")

    p_src.add_run("abc").bold = True
    p_src.add_run("defg").underline = True
    p_src.add_run("hijkl").italic = True
    p_dest = document.add_paragraph("")
    p_dest._p.clear()
    moveTextSegment(p_src, p_dest, 0, 0)
    assert p_src.text == "bcdefghijkl"
    assert p_dest.text == "a"

    p_src._p.clear()
    p_src.add_run("abc").bold = True
    p_src.add_run("defg").underline = True
    p_src.add_run("hijkl").italic = True
    p_dest._p.clear()
    moveTextSegment(p_src, p_dest, 0, 11)
    assert p_src.text == ""
    assert p_dest.text == "abcdefghijkl"

    p_src._p.clear()
    p_src.add_run("abc").bold = True
    p_src.add_run("defg").underline = True
    p_src.add_run("hijkl").italic = True
    p_dest._p.clear()
    moveTextSegment(p_src, p_dest, 3, 6)
    assert p_src.text == "abchijkl"
    assert p_dest.text == "defg"