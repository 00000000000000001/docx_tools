import copy
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def doc_text(doc):
    """
    Kombiniert den äußeren und inneren Text eines Dokumentes und gibt diesen zurück.

    Parameters:
    - param doc: Das Dokument, von dem der Text extrahiert werden soll.
      Dies könnte ein Objekt sein, das sowohl `doc_outer_text` als auch `doc_inner_text` Methoden/Funktionen unterstützt.

    Return:
    - return: Ein String, der die Kombination aus dem äußeren und inneren Text des Dokumentes darstellt.

    Raises:
    - raises AttributeError: Wenn das übergebene `doc`-Objekt nicht die erforderlichen Methoden `doc_outer_text` und `doc_inner_text` unterstützt.
    """
    return doc_outer_text(doc) + doc_inner_text(doc)


def doc_outer_text(doc):
    """
    Extrahiert den äußeren Text eines Dokumentes, indem es den Text jedes Paragraphen sammelt und zusammenfügt.

    Parameters:
    - param doc: Das Dokument, von dem der äußere Text extrahiert werden soll.
      Dies könnte ein Objekt sein, das eine Eigenschaft `paragraphs` hat, welche eine Liste von Paragraph-Objekten ist,
      wobei jedes Paragraph-Objekt eine Eigenschaft `text` hat, die den Text des Paragraphen als String enthält.

    Return:
    - return: Ein String, der den gesamten äußeren Text des Dokumentes darstellt.

    Raises:
    - raises AttributeError: Wenn das übergebene `doc`-Objekt nicht die erforderliche Eigenschaft `paragraphs` besitzt oder
      wenn ein Paragraph-Objekt innerhalb von `doc.paragraphs` nicht die Eigenschaft `text` besitzt.
    """
    fullText = ""
    for p in doc.paragraphs:
        fullText += p.text
    return fullText


def iterate_cells(row, func=lambda x: x.text):
    """
    Iterates over cells in a given row, applying a function to each cell and concatenating the results.

    Parameters:
    - row: The row object containing cells to iterate over.
    - func: A function that takes a cell object as input and returns a string. Defaults to extracting the cell's text.

    Returns:
    - A string that is the concatenation of the function's results for each cell in the row, with a newline character after each cell's output.
    """
    full_text = ""
    for cell in row.cells:
        full_text += func(cell) + "\n"
        if len(cell.tables) > 0:
            full_text += iterate_tables(cell, func)
    return full_text


def iterate_rows(table, func=lambda x: x.text):
    """
    Iterates over rows in a given table, applying the iterate_cells function to each row.

    Parameters:
    - table: The table object containing rows to iterate over.
    - func: A function to be passed to iterate_cells, which is applied to each cell.

    Returns:
    - A string that is the concatenation of all text from all cells in all rows of the table.
    """
    full_text = ""
    for row in table.rows:
        full_text += iterate_cells(row, func)
    return full_text


def iterate_tables(node, func=lambda x: x.text):
    """
    Iterates over tables in a given node (e.g., a document or another table cell), applying the iterate_rows function to each table.

    Parameters:
    - node: The node object containing tables to iterate over.
    - func: A function to be passed to iterate_rows, which is then applied to each cell in each row of each table.

    Returns:
    - A string that is the concatenation of all text from all cells in all tables within the node.
    """
    full_text = ""
    if len(node.tables) > 0:
        for table in node.tables:
            full_text += iterate_rows(table, func)
    return full_text


def doc_inner_text(doc):
    """
    Extrahiert den inneren Text eines Dokumentes durch Iteration über alle Tabellen und sammelt deren Inhalte.

    Parameters:
    - param doc: Das Dokument, von dem der innere Text extrahiert werden soll.
      Das Objekt sollte eine Funktion oder Methode `iterate_tables` unterstützen, die die Iteration durch alle Tabellen im Dokument ermöglicht und deren Inhalte sammelt.

    Return:
    - return: Ein String, der den gesamten inneren Text des Dokumentes darstellt, basierend auf den Inhalten der Tabellen.

    Raises:
    - raises AttributeError: Wenn das übergebene `doc`-Objekt nicht die erforderliche Methode `iterate_tables` besitzt.
    """
    return iterate_tables(doc)


def duplicate(p):
    """
    Dupliziert ein gegebenes Objekt `p` tiefgehend und fügt das duplizierte Objekt in die Sequenz direkt nach `p` ein.

    Parameters:
    - param p: Das Objekt, das dupliziert werden soll.
      Dieses Objekt muss ein Attribut `_p` haben, welches wiederum die Methode `addnext` unterstützen muss, um das duplizierte Objekt in die Sequenz einfügen zu können.

    Return:
    - return: Das duplizierte Objekt `p_new`.

    Raises:
    - raises AttributeError: Wenn das übergebene Objekt `p` nicht das erforderliche Attribut `_p` oder die Methode `addnext` besitzt.
    - raises CopyError: Wenn der tiefe Kopiervorgang (deep copy) fehlschlägt.
    """

    p_new = copy.deepcopy(p)
    p._p.addnext(p_new._p)
    return p_new


def delete_paragraph(paragraph):
    """
    Entfernt einen Absatz aus einem Dokument. Diese Funktion setzt voraus, dass der Absatz ein Objekt mit einer internen Struktur ist,
    die ein `_element`-Attribut besitzt. Das `_element`-Attribut repräsentiert den eigentlichen Absatz im Dokumentenbaum.
    Nach dem Entfernen des Absatzes werden die Referenzen auf das Absatz-Element im übergebenen Absatzobjekt auf `None` gesetzt.

    Parameters:
    - param paragraph: Das Absatzobjekt, das aus seinem Dokument entfernt werden soll.
      Es muss die Attribute `_element` und `_p` haben, wobei `_element` das XML-Element des Absatzes ist und `_p` eine Referenz darauf sein könnte.

    Return:
    - return: Nichts. Die Funktion modifiziert den Zustand des übergebenen Absatzes und seines Elterndokuments direkt.

    Raises:
    - raises AttributeError: Wenn das übergebene Absatzobjekt nicht die erforderlichen Attribute `_element` und `_p` besitzt.
    - raises RemoveError: Wenn das Entfernen des Absatzes aus dem Dokumentenbaum fehlschlägt.
    """
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def append_paragraph(paragraph, text=None, style=None):
    """
    Fügt einen neuen Absatz nach einem gegebenen Absatz hinzu. Optional kann der Text und der Stil des neuen Absatzes spezifiziert werden.

    Parameters:
    - param paragraph: Das Absatzobjekt, nach dem der neue Absatz eingefügt werden soll.
      Es wird erwartet, dass dieses Objekt ein `_p` Attribut für den aktuellen Absatz und ein `_parent` Attribut für das Elternobjekt hat.
    - param text: Optionaler Text, der dem neuen Absatz hinzugefügt wird. Default ist None.
    - param style: Optionaler Stil, der auf den neuen Absatz angewendet wird. Default ist None.

    Return:
    - return: Das Objekt des neu eingefügten Absatzes.

    Raises:
    - raises Exception: Wenn beim Einfügen des neuen Absatzes ein Fehler auftritt, wird eine allgemeine Exception geworfen.
    """
    try:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        return new_para
    except Exception as e:
        print(f"Error when inserting paragraph: {e}")


def delete_paragraph(paragraph):
    """
    Entfernt einen Absatz aus seinem übergeordneten Element im Dokument. Diese Funktion ändert die internen Referenzen des Absatzobjektes,
    indem sie `_element` und `_p` auf `None` setzt, um anzuzeigen, dass der Absatz nicht länger Teil des Dokuments ist.

    Parameters:
    - param paragraph: Das Absatzobjekt, das aus dem Dokument entfernt werden soll.
      Das Objekt muss über ein `_element`-Attribut verfügen, das das XML-Element des Absatzes darstellt.

    Return:
    - return: Nichts. Die Funktion verändert den Zustand des übergebenen Absatzobjektes und entfernt dessen Element aus dem Dokumentenbaum.

    Raises:
    - raises AttributeError: Wenn das übergebene Absatzobjekt nicht die erforderlichen Attribute `_element` oder `_p` besitzt.
    - raises RemoveError: Wenn das Entfernen des Absatzes aus dem Dokumentenbaum fehlschlägt, z.B. weil das Elternelement nicht gefunden werden kann.
    """
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def in_which_run_is(m, p):
    """
    Bestimmt den Index des Textlaufs (`run`), in dem sich das Zeichen an Position 'm' im Text des Absatzes 'p' befindet.

    Parameters:
    - param m: Die Position des Zeichens im Gesamttext des Absatzes, für die der zugehörige Textlauf ermittelt werden soll.
    - param p: Das Absatzobjekt, das die Textläufe (`runs`) enthält. Es wird erwartet, dass dieses Objekt eine Eigenschaft `text` für den Gesamttext
      und eine Liste `runs` für die Textläufe hat.

    Return:
    - return: Der Index des Textlaufs, der das Zeichen an Position 'm' enthält, oder `None`, wenn 'm' außerhalb der Grenzen des Absatztextes liegt
      oder kein entsprechender Textlauf gefunden wird.

    Raises:
    - Es werden keine Exceptions direkt von dieser Funktion ausgelöst, aber sie gibt `None` zurück, wenn die Bedingungen nicht erfüllt sind.
    """
    # Überprüft, ob 'm' außerhalb der Grenzen des Absatztextes liegt.
    if m < 0 or m >= len(p.text):
        return None

    cumulative_length = 0
    # Durchläuft die Textläufe, um sowohl den Index als auch den Textlauf selbst zu erhalten.
    for index, run in enumerate(p.runs):
        cumulative_length += len(run.text)
        # Wenn die kumulative Länge nach Hinzufügen dieses Textlaufs 'm' einschließt, wird der Index zurückgegeben.
        if cumulative_length > m:
            return index

    # Wenn 'm' in keinem der Textläufe gefunden wird, obwohl dies durch die anfängliche Überprüfung abgefangen werden sollte.
    return None


def at_which_position_in_its_run_is(m, p):
    """
    Ermittelt die Position des Zeichens an der Stelle 'm' innerhalb seines Textlaufs im Absatz 'p'.

    Parameters:
    - param m: Die Position des Zeichens im Gesamttext des Absatzes.
    - param p: Das Absatzobjekt, das die Textläufe (`runs`) enthält. Es wird erwartet, dass dieses Objekt eine Liste von Textläufen unter `p.runs` hat.

    Return:
    - return: Die Position des Zeichens innerhalb seines Textlaufs, oder `None`, wenn 'm' außerhalb der Grenzen des Gesamttextes des Absatzes liegt.

    Raises:
    - Es werden keine Exceptions direkt von dieser Funktion ausgelöst, aber sie gibt `None` zurück, wenn die Bedingungen nicht erfüllt sind.
    """
    # Überprüft, ob 'm' innerhalb der gültigen Grenzen des Absatztextes liegt.
    if m < 0 or m >= len(p.text):
        return None

    cumulative_length = (
        0  # Hält die kumulative Länge des Textes bis zum aktuellen Lauf.
    )
    for _, run in enumerate(p.runs):
        previous_cumulative_length = cumulative_length
        cumulative_length += len(run.text)

        # Wenn die kumulative Länge jetzt 'm' einschließt, berechne dessen Position in diesem Lauf.
        if cumulative_length > m:
            # 'm' minus der kumulativen Länge aller vorherigen Läufe ergibt die Position im aktuellen Lauf.
            return m - previous_cumulative_length

    # Für den Fall, dass 'm' irgendwie außerhalb der Gesamttextlänge liegt (was durch die anfängliche Überprüfung abgefangen werden sollte).
    return None


def cp(m, n, p_src, p_dest):
    """
    Kopiert einen Textabschnitt, der durch die Positionen 'm' und 'n' im Quellabsatz `p_src` definiert ist, in den Zielabsatz `p_dest`.

    Parameters:
    - param m: Die Startposition im Quellabsatz `p_src`, ab der der Text kopiert werden soll.
    - param n: Die Endposition im Quellabsatz `p_src`, bis zu der der Text kopiert werden soll.
    - param p_src: Das Quellabsatzobjekt, aus dem der Text kopiert wird. Es wird erwartet, dass dieses Objekt eine Liste von Textläufen (`runs`) enthält.
    - param p_dest: Das Zielabsatzobjekt, zu dem der Text hinzugefügt wird.

    Raises:
    - Es werden keine spezifischen Exceptions direkt von dieser Funktion ausgelöst, aber die Funktion gibt `None` zurück, wenn `m` oder `n` außerhalb der gültigen Grenzen liegen, oder wenn die Start- und Endläufe nicht gefunden werden können.
    """

    # Überprüfung, ob die Start- und Endpositionen innerhalb des gültigen Bereichs des Textes liegen.
    if m < 0 or n > len(p_src.text):
        return
    r_start = in_which_run_is(m, p_src)
    r_finish = in_which_run_is(n, p_src)

    # Überprüfung, ob Start- und Endläufe gefunden wurden.
    if r_start == None or r_finish == None:
        return

    for i in range(r_start, r_finish + 1):
        run = p_src.runs[i]
        r_copy = copy.deepcopy(run)._r

        # Einstellen der Start- und Endposition für den zu kopierenden Text im aktuellen Lauf.
        a = 0
        o = len(run.text)
        if i == r_start:
            a = at_which_position_in_its_run_is(m, p_src)
        if i == r_finish:
            o = at_which_position_in_its_run_is(n, p_src) + 1

        # Setzen des Textes für den kopierten Lauf und Hinzufügen zum Zielabsatz.
        r_copy.text = r_copy.text[a:o]
        p_dest._p.append(r_copy)


def remove_run(run, p):
    """
    Entfernt einen spezifischen Textlauf (`run`) aus einem Absatz (`p`). Diese Funktion durchläuft die Textläufe des Absatzes rückwärts,
    um den zu entfernenden Textlauf zu finden und ihn dann aus dem Absatz zu entfernen.

    Parameters:
    - param run: Das Textlauf-Objekt, das aus dem Absatz entfernt werden soll. Es wird erwartet, dass dieses Objekt ein Attribut `_r` hat,
      welches das zugrundeliegende XML-Element des Textlaufs repräsentiert.
    - param p: Das Absatzobjekt, aus dem der Textlauf entfernt werden soll. Das Objekt sollte eine Liste von Textläufen in `p.runs` und
      ein Attribut `_p` haben, welches das zugrundeliegende XML-Element des Absatzes repräsentiert.

    Return:
    - return: Das Textlauf-Objekt `run`, wenn es gefunden und erfolgreich entfernt wurde. Gibt `None` zurück, wenn der Textlauf im Absatz nicht gefunden wurde.

    Raises:
    - Es werden keine Exceptions direkt von dieser Funktion ausgelöst, aber durch die Verwendung von Attributen wie `_r` und `_p` besteht eine implizite
      Abhängigkeit von der Struktur des Absatz- und Textlaufobjekts, die bei Nichteinhaltung zu Fehlern führen kann.
    """
    i = len(p.runs) - 1
    while i >= 0:
        if p.runs[i]._r == run._r:
            p._p.remove(p.runs[i]._r)
            return run
        i -= 1
    return None


def rm(m, n, p):
    """
    Entfernt Text zwischen den Positionen 'm' und 'n' aus dem Absatz 'p'.

    Parameters:
    - param m: Die Startposition im Text des Absatzes 'p', ab der entfernt werden soll.
    - param n: Die Endposition im Text des Absatzes 'p', bis zu der entfernt werden soll.
    - param p: Das Absatzobjekt, aus dem der Text entfernt wird.

    Return:
    - return: Gibt den modifizierten Absatz 'p' zurück, aus dem der spezifizierte Textbereich entfernt wurde.
              Gibt `None` zurück, falls die Indizes 'm' und 'n' ungültige Werte haben oder die spezifizierten Läufe nicht gefunden werden.

    Raises:
    - Es werden keine spezifischen Exceptions ausgelöst, aber es erfolgt eine Überprüfung auf gültige Indizes und das Vorhandensein der Läufe.
    """
    # Überprüfen der Gültigkeit von 'm' und 'n'
    if m < 0 or m > len(p.text) or n < 0 or n > len(p.text):
        return None

    # Ermitteln der Textläufe, in denen 'm' und 'n' liegen
    r_start = in_which_run_is(m, p)
    r_finish = in_which_run_is(n, p)

    if r_start == None or r_finish == None:
        return None

    # Vorbereitung für das Entfernen des Texts
    a = -1  # Anfangsposition im Startlauf
    o = -1  # Endposition im Endlauf
    arr = []  # Sammlung der zu entfernenden Läufe

    # Durchlaufen der betroffenen Läufe und Bestimmen der genauen Positionen
    for i in range(r_start, r_finish + 1):
        run = p.runs[i]

        if i == r_start:
            a = at_which_position_in_its_run_is(m, p)
        if i == r_finish:
            o = at_which_position_in_its_run_is(n, p)
        if i > r_start and i < r_finish:
            arr.append(run)

    # Entfernen des Texts
    if r_start == r_finish:
        p.runs[r_start].text = p.runs[r_start].text[:a] + p.runs[r_finish].text[o + 1 :]
    else:
        p.runs[r_start].text = p.runs[r_start].text[:a]
        p.runs[r_finish].text = p.runs[r_finish].text[o + 1 :]
    for run in reversed(arr):
        remove_run(run, p)

    return p


def mv(m, n, p_src, p_dest):
    """
    Verschiebt einen Textabschnitt zwischen den Positionen 'm' und 'n' aus dem Quellabsatz `p_src` in den Zielabsatz `p_dest`.

    Parameters:
    - param m: Die Startposition im Text des Quellabsatzes `p_src`, ab der der Text verschoben werden soll.
    - param n: Die Endposition im Text des Quellabsatzes `p_src`, bis zu der der Text verschoben werden soll.
    - param p_src: Das Quellabsatzobjekt, aus dem der Text verschoben wird.
    - param p_dest: Das Zielabsatzobjekt, in das der Text eingefügt wird.

    Raises:
    - Es werden keine spezifischen Exceptions ausgelöst, aber `cp` und `rm` führen jeweils eigene Gültigkeitsprüfungen durch und können in bestimmten Fällen `None` zurückgeben.
    """
    # Kopieren des Textabschnitts von `p_src` nach `p_dest`
    cp(m, n, p_src, p_dest)
    # Entfernen des kopierten Textabschnitts aus `p_src`
    rm(m, n, p_src)


def ins_into_paragraph(str, m, p):
    """
    Fügt einen String 'str' an der Position 'm' in den Absatz 'p' ein.

    Parameters:
    - param str: Der einzufügende String.
    - param m: Die Position im Absatz, an der 'str' eingefügt werden soll.
    - param p: Das Absatzobjekt, in das eingefügt wird. Es wird erwartet, dass dieses Objekt eine Liste von Textläufen (`runs`) hat.

    Return:
    - return: Gibt den modifizierten Absatz 'p' zurück, in den 'str' an der Position 'm' eingefügt wurde.

    Raises:
    - Es werden keine spezifischen Exceptions ausgelöst, aber die Funktion berücksichtigt die Länge der `runs` und passt die Einfügeposition entsprechend an.
    """
    l = 0  # Akkumulierte Länge der Texte in den bisher durchlaufenen Textläufen

    # Fall: Keine Textläufe vorhanden, direktes Einfügen in den Absatztext
    if len(p.runs) == 0:
        p.text = p.text[:m] + str + p.text[m:]
        return p

    # Einfügen in den entsprechenden Textlauf, wenn Textläufe vorhanden sind
    for r in p.runs:
        l += len(r.text)
        if m <= l:
            insert_position = m - (
                l - len(r.text)
            )  # Berechnet die Einfügeposition innerhalb des aktuellen Textlaufs
            r.text = r.text[:insert_position] + str + r.text[insert_position:]
            break

    return p


def replace_text_in_doc(m, p_start, n, p_end, doc, text):
    """
    Ersetzt einen Textabschnitt eines Docx-Dokument durch einen übergebenen String.

    Parameters:
    - param m: Index des Beginns des Textabschnitts im Absatz mit dem Index p_start.
    - param p_start: Index des Absatzen in dem der zu ersetzende Textabschnitt beginnt.
    - param n: Index des Endes des Textabschnitts im Absatz mit dem Index p_end.
    - param p_end: Index des Absatzen in dem der zu ersetzende Textabschnitt endet.
    - param doc: Docx-Dokument in dem der Textabschnitt ersetzt werden soll.
    - param text: Text als String der in das Dokument eingesetzt werden soll.
    """
    if p_start == p_end:
        rm(m, n, doc.paragraphs[p_start])
    else:
        p = doc.paragraphs[p_start]
        rm(m, len(p.text) - 1, p)

        for _ in range(p_start + 1, p_end):
            p = doc.paragraphs[p_start + 1]
            delete_paragraph(p)

        p = doc.paragraphs[p_start + 1]
        rm(0, n, p)

    ins_into_paragraph(text, m, doc.paragraphs[p_start])


def extract_substring_between(m, p_start, n, p_end, docx):
    """
    Extrahiert einen Textbereich aus einem Docx-Dokument, beginnend bei der Position m im Absatz p_start
    und endend bei der Position n im Absatz p_end.

    Parameters:
    - param m: Index, an dem der Text im Absatz p_start beginnt.
    - param p_start: Index des Absatzes, in dem der Text beginnt.
    - param n: Index, an dem der Text im Absatz p_end endet.
    - param p_end: Index des Absatzes, in dem der Text endet.
    - param docx: Das Docx-Dokument, aus dem der Text extrahiert werden soll.

    Return:
    - return: Der extrahierte Text als String.

    Raises:
    - raises Exception: Wenn p_start oder p_end außerhalb der Grenzen der Absätze im Dokument liegen.
    """
    if p_start < 0 or p_start >= len(docx.paragraphs):
        raise ValueError("p_start is out of bounds")
    if p_end < 0 or p_end >= len(docx.paragraphs):
        raise ValueError("p_end is out of bounds")

    if p_start == p_end:
        return docx.paragraphs[p_start].text[m:n]

    string_parts = []
    for i in range(p_start, p_end + 1):
        p_text = docx.paragraphs[i].text
        if i == p_start:
            string_parts.append(p_text[m:])
        elif i == p_end:
            string_parts.append(p_text[:n])
        else:
            string_parts.append(p_text)

    return "".join(string_parts)
