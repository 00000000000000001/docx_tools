
from docx_tools import findPosInRun
from docx import Document

def test_all():
    document = Document()

    p_src = document.add_paragraph("")

    p_src.add_run("abc").bold = True
    p_src.add_run("defg").underline = True
    p_src.add_run("hijkl").italic = True

    assert findPosInRun(0, p_src) == 0
    assert findPosInRun(1, p_src) == 1
    assert findPosInRun(2, p_src) == 2

    assert findPosInRun(3, p_src) == 0
    assert findPosInRun(4, p_src) == 1
    assert findPosInRun(5, p_src) == 2
    assert findPosInRun(6, p_src) == 3

    assert findPosInRun(7, p_src) == 0
    assert findPosInRun(8, p_src) == 1
    assert findPosInRun(9, p_src) == 2
    assert findPosInRun(10, p_src) == 3
    assert findPosInRun(11, p_src) == 4

    assert findPosInRun(12, p_src) == None
    assert findPosInRun(-1, p_src) == None