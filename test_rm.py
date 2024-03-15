from docx_tools import copyTextSegment, removeTextSegment
from docx import Document

def test_rm():
    document = Document()

    p_src = document.add_paragraph("")

    p_src.add_run("abc").bold = True
    p_src.add_run("defg").underline = True
    p_src.add_run("hijkl").italic = True
    p_dest = document.add_paragraph("")

    copyTextSegment(p_src, p_dest, 0, 11)
    removeTextSegment(p_dest, 0, 0)
    assert p_dest.text == "bcdefghijkl"

    p_dest._p.clear()
    copyTextSegment(p_src, p_dest, 0, 11)
    removeTextSegment(p_dest, 5, 7)
    assert p_dest.text == "abcdeijkl"

    p_dest._p.clear()
    copyTextSegment(p_src, p_dest, 0, 11)
    removeTextSegment(p_dest, 0, 2)
    assert p_dest.text == "defghijkl"

    p_dest._p.clear()
    copyTextSegment(p_src, p_dest, 0, 11)
    removeTextSegment(p_dest, 0, 11)
    assert p_dest.text == ""

    p_dest._p.clear()
    copyTextSegment(p_src, p_dest, 0, 11)
    removeTextSegment(p_dest, 1, 1)
    assert p_dest.text == "acdefghijkl"

    p_dest._p.clear()
    copyTextSegment(p_src, p_dest, 0, 11)
    removeTextSegment(p_dest, 1, 12)
    assert p_dest.text == "abcdefghijkl"

    p_dest._p.clear()
    copyTextSegment(p_src, p_dest, 0, 11)
    removeTextSegment(p_dest, 0, 2)
    assert p_dest.text == "defghijkl"

    p_dest._p.clear()
    copyTextSegment(p_src, p_dest, 0, 11)
    removeTextSegment(p_dest, 3, 6)
    assert p_dest.text == "abchijkl"