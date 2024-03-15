from docx import Document
import docx_tools


def test_insert1():
    doc = Document()
    doc.add_paragraph("")
    assert docx_tools.insertStrIntoPara(doc.paragraphs[0], "FOO", 0).text == "FOO"


def test_insert2():
    doc = Document()
    doc.add_paragraph("FOO")
    assert docx_tools.insertStrIntoPara(doc.paragraphs[0], "BAR", 0).text == "BARFOO"
