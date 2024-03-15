from docx_tools import findRunIndex, findPosInRun, copyTextSegment, removeTextSegment, moveTextSegment
from docx import Document

def test_all():
    document = Document()

    p_src = document.add_paragraph("")

    p_src.add_run("abc").bold = True
    p_src.add_run("defg").underline = True
    p_src.add_run("hijkl").italic = True

    assert findRunIndex(0, p_src) == 0
    assert findRunIndex(1, p_src) == 0
    assert findRunIndex(2, p_src) == 0

    assert findRunIndex(3, p_src) == 1
    assert findRunIndex(4, p_src) == 1
    assert findRunIndex(5, p_src) == 1
    assert findRunIndex(6, p_src) == 1

    assert findRunIndex(7, p_src) == 2
    assert findRunIndex(8, p_src) == 2
    assert findRunIndex(9, p_src) == 2
    assert findRunIndex(10, p_src) == 2
    assert findRunIndex(11, p_src) == 2

    assert findRunIndex(12, p_src) == None
    assert findRunIndex(-1, p_src) == None